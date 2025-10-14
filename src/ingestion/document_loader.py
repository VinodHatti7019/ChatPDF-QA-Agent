"""Document loading module for multi-format support.

Supports PDF, DOCX, and TXT file formats with robust error handling
and metadata extraction.
"""

import os
import hashlib
from pathlib import Path
from typing import List, Dict, Optional, Any
from dataclasses import dataclass
from datetime import datetime
import logging

# Document parsers
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import mimetypes

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """Document class to store content and metadata."""
    content: str
    metadata: Dict[str, Any]
    doc_id: str
    source: str
    
    def __post_init__(self):
        """Generate document ID if not provided."""
        if not self.doc_id:
            self.doc_id = self._generate_doc_id()
    
    def _generate_doc_id(self) -> str:
        """Generate unique document ID based on content and source."""
        content_hash = hashlib.md5(
            (self.content + self.source).encode()
        ).hexdigest()
        return f"doc_{content_hash[:16]}"


class DocumentLoader:
    """Universal document loader for multiple file formats."""
    
    SUPPORTED_FORMATS = {
        'pdf': ['.pdf'],
        'docx': ['.docx', '.doc'],
        'txt': ['.txt', '.md', '.text']
    }
    
    def __init__(self):
        """Initialize document loader."""
        self.documents: List[Document] = []
        logger.info("DocumentLoader initialized")
    
    def load(self, file_path: str) -> List[Document]:
        """Load documents from a file path.
        
        Args:
            file_path: Path to the file to load
            
        Returns:
            List of Document objects
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = path.suffix.lower()
        
        # Determine file type
        if file_extension in self.SUPPORTED_FORMATS['pdf']:
            return self._load_pdf(path)
        elif file_extension in self.SUPPORTED_FORMATS['docx']:
            return self._load_docx(path)
        elif file_extension in self.SUPPORTED_FORMATS['txt']:
            return self._load_txt(path)
        else:
            raise ValueError(
                f"Unsupported file format: {file_extension}. "
                f"Supported formats: {self._get_supported_extensions()}"
            )
    
    def load_multiple(self, file_paths: List[str]) -> List[Document]:
        """Load multiple documents.
        
        Args:
            file_paths: List of file paths to load
            
        Returns:
            List of all loaded Document objects
        """
        all_documents = []
        
        for file_path in file_paths:
            try:
                docs = self.load(file_path)
                all_documents.extend(docs)
                logger.info(f"Successfully loaded: {file_path}")
            except Exception as e:
                logger.error(f"Failed to load {file_path}: {str(e)}")
                continue
        
        return all_documents
    
    def _load_pdf(self, path: Path) -> List[Document]:
        """Load PDF document using multiple strategies.
        
        First tries pdfplumber for better text extraction,
        falls back to PyPDF2 if needed.
        """
        documents = []
        
        try:
            # Try pdfplumber first (better extraction)
            with pdfplumber.open(path) as pdf:
                full_text = ""
                metadata = {
                    'source': str(path),
                    'file_type': 'pdf',
                    'num_pages': len(pdf.pages),
                    'loaded_at': datetime.now().isoformat(),
                }
                
                # Extract PDF metadata
                if pdf.metadata:
                    metadata.update({
                        'title': pdf.metadata.get('Title', ''),
                        'author': pdf.metadata.get('Author', ''),
                        'subject': pdf.metadata.get('Subject', ''),
                        'creator': pdf.metadata.get('Creator', ''),
                    })
                
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        full_text += f"\n\n--- Page {page_num} ---\n\n{text}"
                
                if full_text.strip():
                    doc = Document(
                        content=full_text.strip(),
                        metadata=metadata,
                        doc_id="",
                        source=str(path)
                    )
                    documents.append(doc)
                    logger.info(f"Loaded PDF with pdfplumber: {path.name}")
                    return documents
        
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
        
        # Fallback to PyPDF2
        try:
            with open(path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = ""
                metadata = {
                    'source': str(path),
                    'file_type': 'pdf',
                    'num_pages': len(pdf_reader.pages),
                    'loaded_at': datetime.now().isoformat(),
                }
                
                # Extract PDF metadata
                if pdf_reader.metadata:
                    metadata.update({
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                    })
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text:
                        full_text += f"\n\n--- Page {page_num} ---\n\n{text}"
                
                if full_text.strip():
                    doc = Document(
                        content=full_text.strip(),
                        metadata=metadata,
                        doc_id="",
                        source=str(path)
                    )
                    documents.append(doc)
                    logger.info(f"Loaded PDF with PyPDF2: {path.name}")
        
        except Exception as e:
            logger.error(f"Failed to load PDF {path}: {str(e)}")
            raise
        
        return documents
    
    def _load_docx(self, path: Path) -> List[Document]:
        """Load DOCX document."""
        try:
            doc = DocxDocument(path)
            
            # Extract text from paragraphs
            full_text = "\n\n".join(
                [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            )
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        table_text.append(row_text)
            
            if table_text:
                full_text += "\n\n--- Tables ---\n\n" + "\n".join(table_text)
            
            metadata = {
                'source': str(path),
                'file_type': 'docx',
                'num_paragraphs': len(doc.paragraphs),
                'num_tables': len(doc.tables),
                'loaded_at': datetime.now().isoformat(),
            }
            
            # Extract core properties if available
            try:
                core_props = doc.core_properties
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else '',
                })
            except:
                pass
            
            document = Document(
                content=full_text.strip(),
                metadata=metadata,
                doc_id="",
                source=str(path)
            )
            
            logger.info(f"Loaded DOCX: {path.name}")
            return [document]
        
        except Exception as e:
            logger.error(f"Failed to load DOCX {path}: {str(e)}")
            raise
    
    def _load_txt(self, path: Path) -> List[Document]:
        """Load TXT/Markdown document."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(path, 'r', encoding=encoding) as file:
                        content = file.read()
                        used_encoding = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError(f"Could not decode file with any supported encoding")
            
            metadata = {
                'source': str(path),
                'file_type': 'txt',
                'encoding': used_encoding,
                'size_bytes': path.stat().st_size,
                'loaded_at': datetime.now().isoformat(),
            }
            
            document = Document(
                content=content.strip(),
                metadata=metadata,
                doc_id="",
                source=str(path)
            )
            
            logger.info(f"Loaded TXT: {path.name}")
            return [document]
        
        except Exception as e:
            logger.error(f"Failed to load TXT {path}: {str(e)}")
            raise
    
    def _get_supported_extensions(self) -> List[str]:
        """Get list of all supported file extensions."""
        extensions = []
        for format_exts in self.SUPPORTED_FORMATS.values():
            extensions.extend(format_exts)
        return extensions
    
    def get_document_stats(self, documents: List[Document]) -> Dict[str, Any]:
        """Get statistics about loaded documents.
        
        Args:
            documents: List of Document objects
            
        Returns:
            Dictionary containing statistics
        """
        if not documents:
            return {}
        
        total_chars = sum(len(doc.content) for doc in documents)
        file_types = {}
        
        for doc in documents:
            file_type = doc.metadata.get('file_type', 'unknown')
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        return {
            'total_documents': len(documents),
            'total_characters': total_chars,
            'avg_chars_per_doc': total_chars // len(documents),
            'file_types': file_types,
        }


if __name__ == "__main__":
    # Test document loader
    loader = DocumentLoader()
    print("Supported formats:", loader._get_supported_extensions())
